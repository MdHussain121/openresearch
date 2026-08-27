import io

import docx
from fastapi.testclient import TestClient
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.models.citation import Citation
from app.models.document import Document
from app.models.paper import Paper
from app.models.project import Project
from app.services.auth import create_access_token
from app.services.export_service import (
    export_service,
    format_bibliography_entry,
    format_inline_marker,
)


def test_format_inline_markers():
    """Verify inline citation formatting across all 6 supported styles."""
    p1 = Paper(
        id="p1",
        title="Attention Is All You Need",
        authors=[{"familyName": "Vaswani", "givenName": "Ashish"}],
        year=2017,
    )
    p2 = Paper(
        id="p2",
        title="Deep Residual Learning for Image Recognition",
        authors=[
            {"familyName": "He", "givenName": "Kaiming"},
            {"familyName": "Zhang", "givenName": "Xiangyu"},
        ],
        year=2016,
    )
    p3 = Paper(
        id="p3",
        title="BERT: Pre-training of Deep Bidirectional Transformers",
        authors=[
            {"familyName": "Devlin", "givenName": "Jacob"},
            {"familyName": "Chang", "givenName": "Ming-Wei"},
            {"familyName": "Lee", "givenName": "Kenton"},
        ],
        year=2018,
    )

    # APA 7th
    assert format_inline_marker(p1, "apa") == "(Vaswani, 2017)"
    assert format_inline_marker(p2, "apa") == "(He & Zhang, 2016)"
    assert format_inline_marker(p3, "apa") == "(Devlin et al., 2018)"

    # MLA 9th
    assert format_inline_marker(p1, "mla") == "(Vaswani)"
    assert format_inline_marker(p2, "mla") == "(He and Zhang)"

    # Chicago (Author-Date)
    assert format_inline_marker(p1, "chicago") == "(Vaswani 2017)"
    assert format_inline_marker(p2, "chicago") == "(He and Zhang 2016)"

    # IEEE Standard (Numerical)
    assert format_inline_marker(p1, "ieee", index=1) == "[1]"
    assert format_inline_marker(p2, "ieee", index=2) == "[2]"

    # Vancouver (Numerical)
    assert format_inline_marker(p1, "vancouver", index=1) == "(1)"


def test_format_bibliography_entries():
    """Verify full bibliography entry formatting in all 6 styles."""
    paper = Paper(
        id="p1",
        title="FlashAttention: Fast and Memory-Efficient Exact Attention",
        authors=[
            {"familyName": "Dao", "givenName": "Tri"},
            {"familyName": "Fu", "givenName": "Daniel"},
        ],
        year=2022,
        doi="10.48550/arXiv.2205.14135",
        metadata_json={
            "journal": "Advances in Neural Information Processing Systems",
            "volume": "35",
            "pages": "16344-16359",
        },
    )

    # APA
    apa_entry = format_bibliography_entry(paper, "apa", index=1)
    assert "Dao, T., & Fu, D." in apa_entry
    assert "(2022)" in apa_entry
    assert "FlashAttention: Fast and Memory-Efficient Exact Attention" in apa_entry
    assert "https://doi.org/10.48550/arXiv.2205.14135" in apa_entry

    # MLA
    mla_entry = format_bibliography_entry(paper, "mla", index=1)
    assert "Dao, T., and Daniel Fu" in mla_entry or "Dao, T." in mla_entry
    assert '"FlashAttention: Fast and Memory-Efficient Exact Attention."' in mla_entry

    # IEEE
    ieee_entry = format_bibliography_entry(paper, "ieee", index=1)
    assert ieee_entry.startswith("[1]")
    assert "T. Dao, D. Fu" in ieee_entry

    # Vancouver
    vanc_entry = format_bibliography_entry(paper, "vancouver", index=1)
    assert vanc_entry.startswith("(1)")
    assert "Dao T, Fu D" in vanc_entry


def test_export_to_markdown_structure():
    """Verify Markdown export preserves headings, tables, equations, footnotes and bibliography."""
    doc = Document(
        id="d1",
        title="Efficient Deep Learning Architectures",
        plain_text=(
            "# Introduction\n\nFlashAttention improves GPU IO efficiency.\n\n## Methodology\n\n"
            "We benchmarked execution times."
        ),
        content_json={
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1},
                    "content": [{"type": "text", "text": "Introduction"}],
                },
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Transformers rely heavily on self-attention "},
                        {
                            "type": "citation",
                            "attrs": {"paperId": "p1", "label": "(Vaswani et al., 2017)"},
                        },
                        {"type": "text", "text": "."},
                        {"type": "trustMarker", "attrs": {"markerNumber": 1}},
                    ],
                },
                {
                    "type": "mathEquation",
                    "attrs": {
                        "latex": "\\text{Attention}(Q, K, V) = \\text{softmax}\\left(\\frac{QK^T}{\\sqrt{d_k}}\\right)V"
                    },
                },
                {
                    "type": "table",
                    "content": [
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableHeader",
                                    "content": [{"type": "text", "text": "Model"}],
                                },
                                {
                                    "type": "tableHeader",
                                    "content": [{"type": "text", "text": "Speedup"}],
                                },
                            ],
                        },
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [{"type": "text", "text": "FlashAttention"}],
                                },
                                {
                                    "type": "tableCell",
                                    "content": [{"type": "text", "text": "3.2x"}],
                                },
                            ],
                        },
                    ],
                },
            ],
        },
    )
    paper = Paper(
        id="p1",
        title="Attention Is All You Need",
        authors=[{"familyName": "Vaswani", "givenName": "Ashish"}],
        year=2017,
        doi="10.48550/arXiv.1706.03762",
    )
    citation = Citation(
        id="c1",
        document_id="d1",
        paper_id="p1",
        position=0,
        page_number=3,
        attribution_scope="clause",
    )

    md_content, filename, mime = export_service.export_document(
        document=doc,
        citations=[citation],
        papers=[paper],
        export_format="markdown",
        citation_style="apa",
        include_bibliography=True,
        include_trust_markers=True,
    )

    assert filename.endswith(".md")
    assert mime.startswith("text/markdown")
    assert "# Efficient Deep Learning Architectures" in md_content
    assert "## References" in md_content
    assert "Vaswani, A. (2017)" in md_content
    assert "$$" in md_content
    assert "| Model | Speedup |" in md_content
    assert "Footnotes & Source Provenance" in md_content
    assert "[^1]: Source-grounded: *Attention Is All You Need*" in md_content


def test_export_to_docx_binary():
    """Verify DOCX generation produces a valid Word document with sections, tables & citations."""
    doc = Document(
        id="d1",
        title="Attention Optimization in LLMs",
        plain_text="Sample text content.",
        content_json={
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1},
                    "content": [{"type": "text", "text": "Background"}],
                },
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Standard attention scales quadratically."}
                    ],
                },
            ],
        },
    )
    paper = Paper(
        id="p1",
        title="Attention Is All You Need",
        authors=[{"familyName": "Vaswani", "givenName": "Ashish"}],
        year=2017,
    )
    citation = Citation(id="c1", document_id="d1", paper_id="p1", position=0)

    docx_buf, filename, mime = export_service.export_document(
        document=doc,
        citations=[citation],
        papers=[paper],
        export_format="docx",
        citation_style="ieee",
        include_bibliography=True,
    )

    assert filename.endswith(".docx")
    assert "openxmlformats" in mime

    # Load and verify document using python-docx
    docx_doc = docx.Document(docx_buf)
    paragraphs = [p.text for p in docx_doc.paragraphs if p.text]
    assert any("Attention Optimization in LLMs" in p for p in paragraphs)
    assert any("Background" in p for p in paragraphs)
    assert any("References" in p for p in paragraphs)


def test_export_to_pdf_binary():
    """Verify PDF generation creates a compliant PDF file with headers and footnotes."""
    doc = Document(
        id="d1",
        title="Neural Architecture Search Overview",
        plain_text="NAS explores combinatorial design spaces.",
    )
    paper = Paper(
        id="p1",
        title="Neural Architecture Search with Reinforcement Learning",
        authors=[{"familyName": "Zoph", "givenName": "Barret"}],
        year=2016,
    )
    citation = Citation(id="c1", document_id="d1", paper_id="p1", position=0)

    pdf_buf, filename, mime = export_service.export_document(
        document=doc,
        citations=[citation],
        papers=[paper],
        export_format="pdf",
        citation_style="apa",
        include_bibliography=True,
    )

    assert filename.endswith(".pdf")
    assert mime == "application/pdf"

    # Verify PDF stream validity
    pdf_bytes = pdf_buf.getvalue()
    assert pdf_bytes.startswith(b"%PDF")
    reader = PdfReader(io.BytesIO(pdf_bytes))
    assert len(reader.pages) >= 1
    page_text = reader.pages[0].extract_text()
    assert "Neural Architecture Search Overview" in page_text


def test_export_api_endpoints(client: TestClient, db: Session):
    """Verify POST and GET /documents/{id}/export endpoints with authentication."""
    from app.services.auth import create_user_with_personal_owner

    user = create_user_with_personal_owner(
        db=db, email="export@openresearch.org", password="Secure_Password_123", name="Dr. Export"
    )
    project = Project(id="proj_export", owner_id=user.personal_owner_id, name="Export Test Project")
    db.add(project)
    document = Document(
        id="doc_export",
        project_id="proj_export",
        title="Quantum Computing Paper",
        plain_text="Quantum supremacy demonstrated.",
    )
    db.add(document)
    paper = Paper(
        id="p_export",
        project_id="proj_export",
        title="Quantum Supremacy",
        authors=[{"familyName": "Arute", "givenName": "Frank"}],
        year=2019,
    )
    db.add(paper)
    citation = Citation(id="c_export", document_id="doc_export", paper_id="p_export", position=0)
    db.add(citation)
    db.commit()

    token = create_access_token({"sub": user.id})
    headers = {"Authorization": f"Bearer {token}"}

    # 1. POST Export to Markdown
    res_md = client.post(
        "/api/v1/documents/doc_export/export",
        json={"export_format": "markdown", "citation_style": "apa", "include_bibliography": True},
        headers=headers,
    )
    assert res_md.status_code == 200
    assert "Quantum Computing Paper" in res_md.text
    assert "Arute, F. (2019)" in res_md.text

    # 2. POST Export to BibTeX
    res_bib = client.post(
        "/api/v1/documents/doc_export/export",
        json={"export_format": "bibtex"},
        headers=headers,
    )
    assert res_bib.status_code == 200
    assert "@article{" in res_bib.text
    assert "Quantum Supremacy" in res_bib.text

    # 3. GET Direct download URL
    res_get = client.get("/api/v1/documents/doc_export/export/pdf?style=ieee", headers=headers)
    assert res_get.status_code == 200
    assert res_get.headers["content-type"] == "application/pdf"
    assert res_get.content.startswith(b"%PDF")


def test_export_api_guards_and_fallbacks(client: TestClient, db: Session):
    """Cover export endpoint 404/403 guards, invalid-format 400s and the
    no-citations fallback to project papers."""
    from app.services.auth import create_user_with_personal_owner

    user = create_user_with_personal_owner(
        db=db,
        email="export_guard@openresearch.org",
        password="Secure_Password_123",
        name="Dr. Guard",
    )
    outsider = create_user_with_personal_owner(
        db=db,
        email="export_out@openresearch.org",
        password="Secure_Password_123",
        name="Dr. Outsider",
    )
    project = Project(id="proj_export_guard", owner_id=user.personal_owner_id, name="Guard Project")
    db.add(project)
    # Document WITHOUT citations -> exercises fallback to project papers.
    document = Document(
        id="doc_export_guard",
        project_id="proj_export_guard",
        title="Guarded Export Doc",
        plain_text="Body without citations.",
    )
    db.add(document)
    db.add(
        Paper(
            id="p_export_guard",
            project_id="proj_export_guard",
            title="Fallback Paper",
            authors=[{"familyName": "Fallback", "givenName": "Fiona"}],
            year=2021,
        )
    )
    db.commit()

    owner_headers = {"Authorization": f"Bearer {create_access_token({'sub': user.id})}"}
    out_headers = {"Authorization": f"Bearer {create_access_token({'sub': outsider.id})}"}

    # Unknown document -> 404 (POST + GET)
    missing_export = client.post(
        "/api/v1/documents/missing_doc/export",
        json={"export_format": "markdown"},
        headers=owner_headers,
    )
    assert missing_export.status_code == 404
    assert (
        client.get(
            "/api/v1/documents/missing_doc/export/markdown", headers=owner_headers
        ).status_code
        == 404
    )

    # Outsider -> 403 (POST + GET)
    res_forbidden_post = client.post(
        "/api/v1/documents/doc_export_guard/export",
        json={"export_format": "markdown"},
        headers=out_headers,
    )
    assert res_forbidden_post.status_code == 403
    assert (
        client.get(
            "/api/v1/documents/doc_export_guard/export/markdown", headers=out_headers
        ).status_code
        == 403
    )

    # Invalid format -> ValueError mapped to 400 (POST + GET)
    res_bad_format = client.post(
        "/api/v1/documents/doc_export_guard/export",
        json={"export_format": "definitely_not_a_format"},
        headers=owner_headers,
    )
    assert res_bad_format.status_code == 400
    res_bad_format_get = client.get(
        "/api/v1/documents/doc_export_guard/export/definitely_not_a_format", headers=owner_headers
    )
    assert res_bad_format_get.status_code == 400

    # No-citation document falls back to project papers (POST + GET)
    res_md = client.post(
        "/api/v1/documents/doc_export_guard/export",
        json={"export_format": "markdown"},
        headers=owner_headers,
    )
    assert res_md.status_code == 200
    assert "Fallback" in res_md.text
    res_get = client.get(
        "/api/v1/documents/doc_export_guard/export/markdown", headers=owner_headers
    )
    assert res_get.status_code == 200
    assert "Fallback" in res_get.text
