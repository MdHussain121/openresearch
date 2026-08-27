"""
Word Document Exporter (.docx) for OpenResearch Export Engine.
"""

import io
from datetime import UTC, datetime

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models.citation import Citation
from app.models.document import Document
from app.models.paper import Paper
from app.services.export.ast_parser import parse_document_blocks
from app.services.export.csl_formatter import (
    format_authors_inline,
    format_bibliography_entry,
)
from app.services.export.options import ExportOptions


def export_to_docx(
    document: Document,
    citations: list[Citation],
    papers: list[Paper],
    citation_style: str = "apa",
    include_bibliography: bool = True,
    include_trust_markers: bool = True,
    options: ExportOptions | None = None,
) -> io.BytesIO:
    """Generate native Microsoft Word (.docx) file preserving formatting & academic hierarchy."""
    if options is not None:
        citation_style = options.citation_style
        include_bibliography = options.include_bibliography
        include_trust_markers = options.include_trust_markers

    doc = docx.Document()
    paper_dict = {paper.id: paper for paper in papers}

    # Page Margins (1 inch academic standard)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Title (22pt Bold Academic)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(document.title or "Untitled Research Paper")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata subtitle
    sub_p = doc.add_paragraph()
    date_str = datetime.now(UTC).strftime("%B %d, %Y")
    sub_run = sub_p.add_run(
        f"OpenResearch Academic Assistant • Generated on {date_str} • {citation_style.upper()} Style"
    )
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(0x5C, 0x5B, 0x57)
    sub_p.paragraph_format.space_after = Pt(18)

    # Citation map
    citation_map: dict[str, tuple[Paper, int]] = {}
    ordered_papers: list[Paper] = []
    for citation in citations:
        paper = paper_dict.get(citation.paper_id)
        if paper and paper.id not in citation_map:
            ordered_papers.append(paper)
            citation_map[paper.id] = (paper, len(ordered_papers))
    if not ordered_papers and papers:
        for paper in papers:
            ordered_papers.append(paper)
            citation_map[paper.id] = (paper, len(ordered_papers))

    blocks = parse_document_blocks(document, citation_map, citation_style)

    for block in blocks:
        if block.block_type == "heading":
            heading = doc.add_heading(level=min(3, block.level))
            heading_run = heading.runs[0] if heading.runs else heading.add_run()
            heading_run.text = block.content
            heading_run.font.name = "Georgia"
            heading_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)
            if block.level == 1:
                heading_run.font.size = Pt(16)
                heading.paragraph_format.space_before = Pt(14)
                heading.paragraph_format.space_after = Pt(6)
            elif block.level == 2:
                heading_run.font.size = Pt(13)
                heading.paragraph_format.space_before = Pt(10)
                heading.paragraph_format.space_after = Pt(4)
            else:
                heading_run.font.size = Pt(11.5)
                heading.paragraph_format.space_before = Pt(8)
                heading.paragraph_format.space_after = Pt(3)

        elif block.block_type == "paragraph":
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(block.content)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)

        elif block.block_type == "blockquote":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.right_indent = Inches(0.4)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(block.content)
            run.font.name = "Georgia"
            run.font.size = Pt(10.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x3A, 0x39, 0x35)

        elif block.block_type == "code":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.2)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(block.content)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x30)

        elif block.block_type == "equation":
            para = doc.add_paragraph()
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(8)
            run = para.add_run(f"$$ {block.content} $$")
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.font.italic = True

        elif block.block_type == "bullet_list":
            for item in block.children:
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.space_after = Pt(3)
                run = para.add_run(item.content)
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        elif block.block_type == "ordered_list":
            for item in block.children:
                para = doc.add_paragraph(style="List Number")
                para.paragraph_format.space_after = Pt(3)
                run = para.add_run(item.content)
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        elif block.block_type == "table" and block.table_rows:
            num_rows = len(block.table_rows)
            num_cols = max(len(row) for row in block.table_rows) if block.table_rows else 0
            if num_rows > 0 and num_cols > 0:
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = (
                    "Light Shading Accent 1"
                    if "Light Shading Accent 1" in [style.name for style in doc.styles]
                    else "Table Grid"
                )
                for row_index, row in enumerate(block.table_rows):
                    for col_index, cell_value in enumerate(row):
                        cell = table.cell(row_index, col_index)
                        cell.text = cell_value
                        cell_para = cell.paragraphs[0]
                        cell_para.runs[0].font.name = "Calibri"
                        cell_para.runs[0].font.size = Pt(10)
                        if row_index == 0:
                            cell_para.runs[0].font.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Footnotes / Trust Provenance Section
    if include_trust_markers and citations:
        doc.add_heading(level=2).add_run("Footnotes & Source Provenance").font.name = "Georgia"
        for index, citation in enumerate(citations, 1):
            paper = paper_dict.get(citation.paper_id)
            if paper:
                authors_str = format_authors_inline(paper.authors or [], citation_style)
                page_info = f", p. {citation.page_number}" if citation.page_number else ""
                fn_p = doc.add_paragraph()
                fn_p.paragraph_format.space_after = Pt(3)
                fn_num = fn_p.add_run(f"[{index}] ")
                fn_num.font.bold = True
                fn_num.font.size = Pt(9.5)
                fn_body = fn_p.add_run(
                    f"Source-grounded: {paper.title}, {authors_str} ({paper.year or 'n.d.'}){page_info}"
                )
                fn_body.font.name = "Calibri"
                fn_body.font.size = Pt(9.5)
                fn_body.font.color.rgb = RGBColor(0x5C, 0x5B, 0x57)

    # References Section
    if include_bibliography and ordered_papers:
        doc.add_heading(level=1).add_run("References").font.name = "Georgia"
        for index, paper in enumerate(ordered_papers, 1):
            ref_entry = format_bibliography_entry(paper, citation_style, index)
            ref_p = doc.add_paragraph()
            ref_p.paragraph_format.space_after = Pt(4)
            ref_p.paragraph_format.left_indent = Inches(0.3)
            ref_p.paragraph_format.first_line_indent = Inches(-0.3)  # Hanging indent
            ref_run = ref_p.add_run(ref_entry)
            ref_run.font.name = "Calibri"
            ref_run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
